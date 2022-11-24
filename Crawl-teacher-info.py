#8822-8729
#https://it.ctgu.edu.cn/info/1168/8822.htm
#主讲课程</span>
#成果</span>
from docx import Document
import requests
import re

# 根据url获取网页html内容
def getHtml(url,headers):
  page = requests.get(url,headers)
  if page.status_code == 404:
    return
  page.encoding = 'utf-8'
  return page.text

def getInfo(html_s):
  Info=[]
  pattern_a = re.compile('<HTML><HEAD><TITLE>(.+?)-三峡大学计算机与信息学院')
  results =  re.search(pattern_a,html_s)
  # print(results.group(1))
  Info.append(results.group(1))
  pattern_a = re.compile('<!--StartFragment-->((?:.|\n)*?)<!--EndFragment-->')
  html =  re.search(pattern_a,html_s)
  if not(html):
    pattern_a = re.compile('<!---top---->((?:.|\n)*?)<!---last---->')
    html =  re.search(pattern_a,html_s)
  pattern_a = re.compile('<.+?>|&nbsp;|\r{2,}')
  html = re.sub(pattern_a,'',html.group(1))
  Info.append(html)
  # results = re.search('个人基本情况((?:.|\n)*?)主要研究方向',html)
  # Info.append(results.group(1))
  # results = re.search('主要研究方向((?:.|\n)*?)主讲课程',html)
  # Info.append(results.group(1))
  # results = re.search('主讲课程((?:.|\n)*?)近年主要教科研成果',html)
  # Info.append(results.group(1))
  # results = re.search('近年主要教科研成果((?:.|\n)*?)邮箱',html)
  # Info.append(results.group(1))
  # Info.append(results.group(1))
  return Info

def save_docx(Info):
  document = Document('C:/Users/HP/Desktop/实例.docx')
  document.add_heading(Info[0], 1)
  # document.add_heading('个人基本情况', level=3)
  document.add_paragraph(Info[1])
  # document.add_heading('主要研究方向', level=3)
  # document.add_paragraph(Info[2])
  # document.add_heading('主讲课程', level=3)
  # document.add_paragraph(Info[3])
  # document.add_heading('近年主要教科研成果', level=3)
  # document.add_paragraph(Info[4])
  document.save('C:/Users/HP/Desktop/实例.docx') #保存文件

def download(url,headers):
  html = getHtml(url,headers)
  if html:
    Info = getInfo(html)
  # print (Info)
    save_docx(Info)
  else:
    return
  # save_docx(html)
  
def main():
  headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                      'Chrome/91.0.4472.124 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,'
                  'application/signed-exchange;v=b3;q=0.9',
    }
  url_s = 'https://it.ctgu.edu.cn/info/1168/'
  for i in range(8729,8822+1):
    url = url_s + str(i) + ".htm"
    download(url,headers)
    print('fininsh:'+str(i))
  
if __name__ == '__main__':
  main()